"""
skills/plugins/data_summarize_doc.py — Data: Summarize Document

Extracts and returns text from PDF, DOCX, or plain text files.
No LLM call — pure extraction with structural metadata.

Risk: LOW — fs:read
"""
from __future__ import annotations
import asyncio, time
from pathlib import Path
from typing import ClassVar
from neuralclaw.skills.base import SkillBase
from neuralclaw.skills.types import RiskLevel, SkillManifest, SkillResult, SkillValidationError

class DataSummarizeDocSkill(SkillBase):
    manifest: ClassVar[SkillManifest] = SkillManifest(
        name="data_summarize_doc",
        version="1.0.0",
        description="Extract text content from PDF, DOCX, or plain text files. Returns extracted text, page/word count, and metadata. Requires pypdf for PDF, python-docx for DOCX.",
        category="data",
        risk_level=RiskLevel.LOW,
        capabilities=frozenset({"fs:read"}),
        timeout_seconds=30,
        parameters={"type":"object","properties":{
            "file_path":{"type":"string","description":"Path to document file (PDF, DOCX, TXT, MD)."},
            "max_chars":{"type":"integer","description":"Max characters of text to return (default 10000).","default":10000},
            "pages":{"type":"array","items":{"type":"integer"},"description":"Specific pages to extract (1-based, PDF only). Default: all.","default":[]},
        },"required":["file_path"]},
    )

    async def validate(self, file_path: str, **_) -> None:
        p = Path(file_path).expanduser()
        if not p.exists(): raise SkillValidationError(f"File does not exist: '{file_path}'")
        if p.suffix.lower() not in (".pdf",".docx",".txt",".md",".rst",".html"):
            raise SkillValidationError(f"Unsupported file type: '{p.suffix}'. Supported: .pdf .docx .txt .md .rst .html")

    async def execute(self, file_path: str, max_chars: int=10000, pages: list|None=None, **kwargs) -> SkillResult:
        call_id = kwargs.get("_skill_call_id","")
        t_start = time.monotonic()
        max_chars = min(int(max_chars), 100_000)
        p = Path(file_path).expanduser().resolve()

        def _extract():
            suffix = p.suffix.lower()
            if suffix == ".pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(str(p))
                    total_pages = len(reader.pages)
                    page_nums = [i-1 for i in (pages or range(1, total_pages+1)) if 0 < i <= total_pages]
                    text = "\n".join(reader.pages[i].extract_text() or "" for i in page_nums)
                    return {"format":"pdf","pages_total":total_pages,"pages_extracted":len(page_nums),
                            "text":text[:max_chars],"word_count":len(text.split())}
                except ImportError:
                    return {"error":"pypdf not installed. Run: pip install pypdf","format":"pdf"}
            elif suffix == ".docx":
                try:
                    import docx
                    doc = docx.Document(str(p))
                    text = "\n".join(p.text for p in doc.paragraphs)
                    return {"format":"docx","paragraph_count":len(doc.paragraphs),
                            "text":text[:max_chars],"word_count":len(text.split())}
                except ImportError:
                    return {"error":"python-docx not installed. Run: pip install python-docx","format":"docx"}
            else:
                text = p.read_text(errors="replace")
                return {"format":suffix.lstrip("."),"text":text[:max_chars],"word_count":len(text.split()),"char_count":len(text)}

        try:
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, _extract)
            result["file"] = str(p)
            duration_ms = (time.monotonic()-t_start)*1000
            return SkillResult.ok(self.manifest.name, call_id, result, duration_ms=duration_ms)
        except BaseException as e:
            return SkillResult.fail(self.manifest.name, call_id, f"{type(e).__name__}: {e}", type(e).__name__,
                                    duration_ms=(time.monotonic()-t_start)*1000)
